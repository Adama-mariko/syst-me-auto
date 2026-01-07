import io
from reportlab.lib import colors
from reportlab.lib.pagesizes import letter, landscape, A4
from reportlab.platypus import SimpleDocTemplate, Table, TableStyle, Paragraph, Spacer
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import cm
import openpyxl
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.shared import RGBColor
from datetime import datetime

class ReportService:
    @staticmethod
    def generate_word_report(requests, title="Rapport Global des Requêtes"):
        doc = Document()
        
        # Titre
        heading = doc.add_heading(title, 0)
        heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Sous-titre
        subtitle = doc.add_paragraph(f"Généré le : {datetime.now().strftime('%d/%m/%Y à %H:%M')}")
        subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
        doc.add_paragraph() # Espace
        
        # Tableau
        # Colonnes: ID, Date, Type, Prio, Demandeur, Description, Statut, Agent, Echéance
        table = doc.add_table(rows=1, cols=9)
        table.style = 'Table Grid'
        
        # En-têtes
        hdr_cells = table.rows[0].cells
        headers = ['ID', 'Date', 'Type', 'Priorité', 'Demandeur', 'Description', 'Statut', 'Agent', 'Echéance']
        
        for i, header in enumerate(headers):
            cell = hdr_cells[i]
            cell.text = header
            # Style gras pour l'en-tête
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.bold = True
                    run.font.color.rgb = RGBColor(255, 255, 255)
            # Fond bleu foncé (nécessite manipulation XML ou juste accepter le style par défaut, ici on reste simple ou on utilise un style Word)
            # Pour simplifier avec python-docx sans XML complexe, on s'appuie sur le style 'Table Grid' mais on ne peut pas facilement mettre de fond de couleur sans XML.
            # On va laisser le style par défaut 'Table Grid' qui est propre.
            
        for req in requests:
            row_cells = table.add_row().cells
            
            # ID
            row_cells[0].text = str(req.id)
            
            # Date
            row_cells[1].text = req.created_at.strftime('%d/%m/%Y') if req.created_at else ''
            
            # Type
            row_cells[2].text = req.request_type
            
            # Priorité
            row_cells[3].text = req.priority
            
            # Demandeur (Nom + Service + Contact)
            row_cells[4].text = f"{req.requester_name}\n{req.requester_service}\n{req.requester_contact}"
            
            # Description
            row_cells[5].text = req.description
            
            # Statut
            row_cells[6].text = req.status
            
            # Agent
            row_cells[7].text = req.assigned_agent or 'Non assigné'
            
            # Echéance (date de résolution si disponible, sinon deadline)
            if req.resolution_date:
                row_cells[8].text = req.resolution_date.strftime('%d/%m/%Y')
            else:
                row_cells[8].text = req.deadline.strftime('%d/%m/%Y') if req.deadline else '-'
            
            # Ajustement de la taille de police pour tout le tableau
            for cell in row_cells:
                for paragraph in cell.paragraphs:
                    paragraph.style.font.size = Pt(8)

        # Orientation Paysage (car beaucoup de colonnes)
        section = doc.sections[0]
        new_width, new_height = section.page_height, section.page_width
        section.orientation = WD_ALIGN_PARAGRAPH.CENTER # Hacky constant, but for orientation enum is different.
        # Correct way for landscape:
        section.page_width = new_width
        section.page_height = new_height

        buffer = io.BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        return buffer

    @staticmethod
    def generate_pdf_report(requests, title="Rapport Global des Requêtes"):
        buffer = io.BytesIO()
        doc = SimpleDocTemplate(buffer, pagesize=landscape(A4))
        elements = []
        styles = getSampleStyleSheet()
        title_style = ParagraphStyle('CustomTitle', parent=styles['Title'], fontSize=24, spaceAfter=20, textColor=colors.HexColor('#2c3e50'))
        elements.append(Paragraph(title, title_style))
        subtitle = f"Généré le : {datetime.now().strftime('%d/%m/%Y à %H:%M')}"
        elements.append(Paragraph(subtitle, styles['Normal']))
        elements.append(Spacer(1, 20))
        cell_style = ParagraphStyle('CellStyle', parent=styles['Normal'], fontSize=8, leading=10)
        header_style = ParagraphStyle('HeaderStyle', parent=styles['Normal'], fontSize=10, leading=12, textColor=colors.whitesmoke, fontName='Helvetica-Bold')
        headers = [
            Paragraph('ID', header_style),
            Paragraph('Date', header_style),
            Paragraph('Type', header_style),
            Paragraph('Priorité', header_style),
            Paragraph('Demandeur', header_style),
            Paragraph('Description', header_style),
            Paragraph('Echéance', header_style),
            Paragraph('Statut', header_style),
            Paragraph('Agent', header_style)
        ]
        data = [headers]
        for req in requests:
            requester_info = f"<b>{req.requester_name}</b><br/>{req.requester_service}<br/>{req.requester_contact}"
            deadline_str = (
                req.resolution_date.strftime('%d/%m/%Y')
                if req.resolution_date
                else (req.deadline.strftime('%d/%m/%Y') if req.deadline else '-')
            )
            desc = req.description or ''
            if len(desc) > 100:
                desc = desc[:100] + "..."
            row = [
                str(req.id),
                req.created_at.strftime('%d/%m/%Y') if req.created_at else '',
                Paragraph(req.request_type, cell_style),
                req.priority,
                Paragraph(requester_info, cell_style),
                Paragraph(desc, cell_style),
                deadline_str,
                req.status,
                Paragraph(req.assigned_agent or 'Non assigné', cell_style)
            ]
            data.append(row)
        col_widths = [30, 60, 80, 60, 110, 220, 60, 70, 80]
        table = Table(data, colWidths=col_widths, repeatRows=1)
        table.setStyle(TableStyle([
            ('BACKGROUND', (0, 0), (-1, 0), colors.HexColor('#2c3e50')),
            ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
            ('ALIGN', (0, 0), (-1, -1), 'LEFT'),
            ('VALIGN', (0, 0), (-1, -1), 'TOP'),
            ('GRID', (0, 0), (-1, -1), 0.5, colors.black),
            ('BOTTOMPADDING', (0, 0), (-1, -1), 6),
            ('TOPPADDING', (0, 0), (-1, -1), 6),
        ]))
        elements.append(table)
        elements.append(Spacer(1, 30))
        elements.append(Paragraph("Rapport généré automatiquement par le Système Auto.", styles['Italic']))
        doc.build(elements)
        buffer.seek(0)
        return buffer

    @staticmethod
    def generate_excel_report(requests):
        wb = openpyxl.Workbook()
        ws = wb.active
        ws.title = "Requêtes Complètes"
        
        # Style des en-têtes
        headers = [
            'ID', 'Date Création', 'Type', 'Priorité', 
            'Demandeur - Nom', 'Demandeur - Service', 'Demandeur - Contact',
            'Description', 'Statut', 'Date Résolution', 
            'Agent Assigné', 'Echéance', 'Commentaires Internes'
        ]
        ws.append(headers)
        
        # Mettre en gras les en-têtes
        for cell in ws[1]:
            cell.font = openpyxl.styles.Font(bold=True)
            cell.fill = openpyxl.styles.PatternFill(start_color="CCE5FF", end_color="CCE5FF", fill_type="solid")
        
        for req in requests:
            ws.append([
                req.id,
                req.created_at,
                req.request_type,
                req.priority,
                req.requester_name,
                req.requester_service,
                req.requester_contact,
                req.description,
                req.status,
                req.resolution_date,
                req.assigned_agent,
                req.resolution_date if req.resolution_date else req.deadline,
                req.internal_comments
            ])
            
        # Ajustement automatique des colonnes
        for col in ws.columns:
            max_length = 0
            column = col[0].column_letter # Get the column name
            for cell in col:
                try:
                    if len(str(cell.value)) > max_length:
                        max_length = len(str(cell.value))
                except:
                    pass
            adjusted_width = (max_length + 2)
            ws.column_dimensions[column].width = min(adjusted_width, 50) # Max 50 de large

        buffer = io.BytesIO()
        wb.save(buffer)
        buffer.seek(0)
        return buffer
